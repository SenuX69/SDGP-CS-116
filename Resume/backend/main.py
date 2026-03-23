from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import os, re, tempfile, shutil
from pathlib import Path

import pdfplumber
from docx import Document

app = FastAPI(title="Resume Scanner API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = Path("./outputs")
OUTPUT_DIR.mkdir(exist_ok=True)


# Extracting Text

def extract_pdf(path: str) -> str:
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text.strip()


def extract_docx(path: str) -> str:
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)


# Section Detection

EMAIL_RE    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_RE    = re.compile(r"(\+?[\d][\d\s\-().]{7,}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[^\s,]+", re.I)
GITHUB_RE   = re.compile(r"github\.com/[^\s,]+", re.I)
URL_RE      = re.compile(r"https?://[^\s,]+", re.I)

SECTIONS = {
    "summary":        ["summary", "objective", "profile", "about me", "about", "overview", "career objective"],
    "experience":     ["experience", "work experience", "employment history", "work history", "professional experience"],
    "education":      ["education", "academic background", "qualifications", "schooling"],
    "skills":         ["skills", "technical skills", "core competencies", "technologies", "key skills", "competencies"],
    "projects":       ["projects", "personal projects", "key projects", "portfolio"],
    "certifications": ["certifications", "certificates", "licenses", "credentials"],
    "languages":      ["languages", "language proficiency", "spoken languages"],
    "interests":      ["interests", "hobbies", "activities", "extracurricular"],
    "references":     ["references", "referees"],
    "achievements":   ["achievements", "awards", "honors", "accomplishments"],
}

# Skill names

SKILL_CANONICAL: dict[str, str] = {
    "node.js": "Node.js",
    "node":    "Node.js",
    "react native": "React Native",
    "react":   "React",
    "vue":     "Vue.js",
    "svelte":  "Svelte",
    "angular": "Angular",
    "next.js": "Next.js",
    "express": "Express",
    "django":  "Django",
    "flask":   "Flask",
    "fastapi": "FastAPI",
    "spring":  "Spring",
    "laravel": "Laravel",
    "python":  "Python",
    "java":    "Java",
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "kotlin":  "Kotlin",
    "swift":   "Swift",
    "ruby":    "Ruby",
    "php":     "PHP",
    "c++":     "C++",
    "c#":      "C#",
    ".net":    ".NET",
    "go":      "Go",
    "rust":    "Rust",
    "dart":    "Dart",
    "flutter": "Flutter",
    "sql":     "SQL",
    "mysql":   "MySQL",
    "postgresql": "PostgreSQL",
    "mongodb": "MongoDB",
    "redis":   "Redis",
    "sqlite":  "SQLite",
    "oracle":  "Oracle",
    "cassandra": "Cassandra",
    "elasticsearch": "Elasticsearch",
    "kafka":   "Kafka",
    "rabbitmq": "RabbitMQ",
    "docker":  "Docker",
    "kubernetes": "Kubernetes",
    "aws":     "AWS",
    "azure":   "Azure",
    "gcp":     "GCP",
    "terraform": "Terraform",
    "ansible": "Ansible",
    "jenkins": "Jenkins",
    "ci/cd":   "CI/CD",
    "git":     "Git",
    "github":  "GitHub",
    "gitlab":  "GitLab",
    "html":    "HTML",
    "css":     "CSS",
    "sass":    "Sass",
    "tailwind": "Tailwind CSS",
    "bootstrap": "Bootstrap",
    "rest":    "REST API",
    "graphql": "GraphQL",
    "grpc":    "gRPC",
    "linux":   "Linux",
    "bash":    "Bash",
    "shell":   "Shell",
    "tensorflow": "TensorFlow",
    "pytorch": "PyTorch",
    "scikit-learn": "Scikit-learn",
    "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "nlp":     "NLP",
    "data analysis": "Data Analysis",
    "pandas":  "Pandas",
    "numpy":   "NumPy",
    "excel":   "Excel",
    "tableau": "Tableau",
    "power bi": "Power BI",
    "figma":   "Figma",
    "photoshop": "Photoshop",
    "illustrator": "Illustrator",
    "agile":   "Agile",
    "scrum":   "Scrum",
    "jira":    "Jira",
    "firebase": "Firebase",
    "supabase": "Supabase",
    "nginx":   "Nginx",
}

# Skills considered to match with roles
REQUIRED_SKILLS = {
    "Python", "JavaScript", "TypeScript", "SQL", "Git",
    "Docker", "React", "Node.js", "AWS", "REST API",
    "Linux", "Agile", "HTML", "CSS",
}


def detect_sections(lines: list[str]) -> dict[str, int]:
    found: dict[str, int] = {}
    for i, line in enumerate(lines):
        norm = line.lower().strip().rstrip(":").strip()
        for sec, keywords in SECTIONS.items():
            if norm in keywords and sec not in found:
                found[sec] = i
    return found


def get_section_text(lines: list[str], smap: dict[str, int], key: str) -> str:
    if key not in smap:
        return ""
    start = smap[key] + 1
    nexts = sorted(v for k, v in smap.items() if k != key and v > start)
    end   = nexts[0] if nexts else len(lines)
    return "\n".join(lines[start:end]).strip()


# Personal Details Extraction 

def extract_name(lines: list[str]) -> str:
    for line in lines[:6]:
        s = line.strip()
        if (
            s and 2 <= len(s.split()) <= 5
            and not EMAIL_RE.search(s)
            and not PHONE_RE.search(s)
            and not any(s.lower().startswith(kw) for kws in SECTIONS.values() for kw in kws)
            and not any(c.isdigit() for c in s)
        ):
            return s
    return ""


# Skill Extraction & Normalization

def extract_skills(text: str) -> list[str]:

    # Matching skills from SKILL_CANONICAL against the text. Returns deduplicated, display-name list.
    
    lower = text.lower()
    found_display: set[str] = set()

    # Sort patterns longest-first to give multi-word patterns priority
    for pattern in sorted(SKILL_CANONICAL, key=len, reverse=True):
        display = SKILL_CANONICAL[pattern]
        if display in found_display:
            continue 
        if re.search(r"\b" + re.escape(pattern) + r"\b", lower):
            found_display.add(display)

    return sorted(found_display)


def rank_skills(skills: list[str]) -> list[str]:

    # Return skills sorted: commonly required first, then the rest.
    
    required = sorted(s for s in skills if s in REQUIRED_SKILLS)
    rest = sorted(s for s in skills if s not in REQUIRED_SKILLS)
    return required + rest

# Experience Parsing 

# Patterns that hint at a duration / date range like "Jan 2021 – Dec 2022" or "2019 - Present"
DURATION_RE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*\d{4}"
    r"|(?:19|20)\d{2})"
    r"\s*[-–—to]+\s*"
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]*\d{4}"
    r"|present|current|now|(?:19|20)\d{2})",
    re.I,
)

# Typical role title keywords used to detect a new entry boundary
ROLE_KEYWORDS = re.compile(
    r"\b(engineer|developer|manager|analyst|designer|architect|consultant|"
    r"lead|director|intern|specialist|scientist|administrator|coordinator)\b",
    re.I,
)


def parse_experience_entries(text: str) -> list[dict]:
    """
    Attempt to split a blob of experience text into structured entries.
    Falls back gracefully — if we can't detect structure, returns one
    entry with the raw text as the description.
    """
    if not text:
        return []

    lines  = [l.strip() for l in text.splitlines() if l.strip()]
    entries: list[dict] = []
    current: dict | None = None

    for line in lines:
        duration_match = DURATION_RE.search(line)
        is_role_line   = bool(ROLE_KEYWORDS.search(line)) and len(line) < 120

        if is_role_line or duration_match:
            # Save the previous entry
            if current:
                entries.append(current)

            duration = duration_match.group(0).strip() if duration_match else ""
            # Removes the duration from the line to leave company/role text
            role_line = DURATION_RE.sub("", line).strip(" -–—|,")

            # If there's a separator character, split company from role
            for sep in [" at ", " @ ", " | ", " — ", " – ", " - "]:
                if sep.lower() in role_line.lower():
                    parts = re.split(re.escape(sep), role_line, maxsplit=1, flags=re.I)
                    current = {
                        "role":        parts[0].strip(),
                        "company":     parts[1].strip(),
                        "duration":    duration,
                        "description": "",
                    }
                    break
            else:
                current = {
                    "role":        "",
                    "company":     role_line,
                    "duration":    duration,
                    "description": "",
                }
        else:
            # Accumulate description lines
            if current is None:
                current = {"role": "", "company": "", "duration": "", "description": ""}
            sep = "\n" if current["description"] else ""
            current["description"] = current["description"] + sep + line

    if current:
        entries.append(current)

    # If only a single entry is avalable and it has no role or company the text is unstructured —
    # return it as a single raw-description entry rather instead of parsing it.
    if len(entries) == 1 and not entries[0]["role"] and not entries[0]["company"]:
        return [{"role": "", "company": "", "duration": "", "description": text}]

    return entries


# Education Parsing

DEGREE_RE = re.compile(
    r"\b(b\.?sc|b\.?tech|b\.?e|b\.?a|m\.?sc|m\.?tech|m\.?b\.?a|ph\.?d|"
    r"bachelor|master|doctorate|diploma|associate|hnd)\b",
    re.I,
)


def parse_education_entries(text: str) -> list[dict]:
    """
    Attempt to split education text into structured entries.
    """
    if not text:
        return []

    lines   = [l.strip() for l in text.splitlines() if l.strip()]
    entries: list[dict] = []
    current: dict | None = None

    for line in lines:
        duration_match = DURATION_RE.search(line)
        degree_match   = DEGREE_RE.search(line)

        if degree_match or (duration_match and len(line) < 120):
            if current:
                entries.append(current)
            duration    = duration_match.group(0).strip() if duration_match else ""
            clean_line  = DURATION_RE.sub("", line).strip(" -–—|,")
            current = {
                "institution": "",
                "degree":      clean_line,
                "duration":    duration,
            }
        else:
            if current is None:
                current = {"institution": "", "degree": "", "duration": ""}
            # First unparsed line after a degree line is likely the institution or else ignoring additional detail lines such as (grades,notes, ect)
            if current["institution"] == "" and line:
                current["institution"] = line

    if current:
        entries.append(current)

    if not entries:
        return [{"institution": text, "degree": "", "duration": ""}]

    return entries


# Resume Analyzer

def analyze_resume(parsed: dict) -> dict:
    
    #Analyzing resume and producing actionable suggestions for users cv's improvement
    
    skills_set = {s for s in parsed.get("skills", [])}

    matched_required = skills_set & REQUIRED_SKILLS
    missing_required = REQUIRED_SKILLS - skills_set
    info = parsed.get("personal_info", {})

    # Suggestions
    suggestions: list[str] = []

    if missing_required:
        top_missing = sorted(missing_required)[:5]
        suggestions.append(f"Consider adding these commonly required skills: {', '.join(top_missing)}.")

    if not parsed.get("summary"):
        suggestions.append("Add a professional summary or objective at the top of your resume.")

    if not info.get("links"):
        suggestions.append("Include a LinkedIn profile or GitHub URL to strengthen your online presence.")

    if not parsed.get("projects"):
        suggestions.append("Add a Projects section to showcase hands-on work.")

    if not parsed.get("certifications"):
        suggestions.append("Relevant certifications (AWS, GCP, Azure, etc.) can significantly boost your profile.")

    if len(skills_set) < 8:
        suggestions.append("Expand your Skills section — aim for at least 8–12 relevant technologies.")

    experience_entries = parsed.get("experience", [])
    if isinstance(experience_entries, list):
        for entry in experience_entries:
            if not entry.get("duration"):
                suggestions.append("Make sure each work experience entry includes clear date ranges.")
                break

    return {
        "matched_skills":  sorted(matched_required),
        "missing_skills":  sorted(missing_required),
        "suggestions": suggestions,
    }


# Core Parser

def parse_resume(raw: str) -> dict:
    lines = [l for l in raw.splitlines() if l.strip()]
    smap  = detect_sections(lines)

    emails = EMAIL_RE.findall(raw)
    phones = PHONE_RE.findall(raw)
    links  = list(dict.fromkeys(
        LINKEDIN_RE.findall(raw) + GITHUB_RE.findall(raw) + URL_RE.findall(raw)
    ))

    skills_text = get_section_text(lines, smap, "skills") or raw
    raw_skills  = extract_skills(skills_text)
    ranked      = rank_skills(raw_skills)

    exp_text = get_section_text(lines, smap, "experience")
    edu_text = get_section_text(lines, smap, "education")

    return {
        "personal_info": {
            "name":  extract_name(lines),
            "email": emails[0] if emails else "",
            "phone": phones[0].strip() if phones else "",
            "links": links,
        },
        "summary":        get_section_text(lines, smap, "summary"),
        "skills":         ranked,
        "experience":     parse_experience_entries(exp_text),
        "education":      parse_education_entries(edu_text),
        "projects":       get_section_text(lines, smap, "projects"),
        "certifications": get_section_text(lines, smap, "certifications"),
        "achievements":   get_section_text(lines, smap, "achievements"),
        "languages":      get_section_text(lines, smap, "languages"),
        "interests":      get_section_text(lines, smap, "interests"),
        "references":     get_section_text(lines, smap, "references"),
    }


# Text Preview Formatter

def format_txt(data: dict, analysis: dict, filename: str) -> str:
    SEP  = "=" * 62
    SEP2 = "-" * 42
    out  = []
    out.append(SEP)
    out.append("               RESUME SCAN RESULTS")
    out.append(f"  Source : {filename}")
    out.append(SEP)

    p = data["personal_info"]
    out.append("\n[ PERSONAL DETAILS ]")
    out.append(SEP2)
    out.append(f"  Name   : {p['name']  or 'N/A'}")
    out.append(f"  Email  : {p['email'] or 'N/A'}")
    out.append(f"  Phone  : {p['phone'] or 'N/A'}")
    if p["links"]:
        out.append("  Links  :")
        for lnk in p["links"]:
            out.append(f"           {lnk}")

    def block(title: str, content):
        if not content:
            return
        out.append(f"\n[ {title.upper()} ]")
        out.append(SEP2)
        if isinstance(content, list):
            # List of dicts (experience / education entries)
            if content and isinstance(content[0], dict):
                for entry in content:
                    parts = [v for v in [entry.get("role"), entry.get("company"), entry.get("duration")] if v]
                    if parts:
                        out.append(f"  >> {' | '.join(parts)}")
                    if entry.get("description"):
                        for ln in entry["description"].strip().splitlines():
                            out.append(f"     {ln}")
                    if entry.get("degree") or entry.get("institution"):
                        deg  = entry.get("degree", "")
                        inst = entry.get("institution", "")
                        dur  = entry.get("duration", "")
                        label = " | ".join(x for x in [deg, inst, dur] if x)
                        out.append(f"  >> {label}")
            else:
                for item in content:
                    out.append(f"  * {item}")
        else:
            for ln in str(content).strip().splitlines():
                out.append(f"  {ln}")

    block("Summary / Objective", data["summary"])
    block("Skills",              data["skills"])
    block("Work Experience",     data["experience"])
    block("Education",           data["education"])
    block("Projects",            data["projects"])
    block("Certifications",      data["certifications"])
    block("Achievements",        data["achievements"])
    block("Languages",           data["languages"])
    block("Interests / Hobbies", data["interests"])
    block("References",          data["references"])

    # Analysis summary
    out.append(f"\n[ RESUME ANALYSIS ]")
    out.append(SEP2)
    out.append(f"  Matched Skills : {', '.join(analysis['matched_skills']) or 'None'}")
    out.append(f"  Missing Skills : {', '.join(analysis['missing_skills']) or 'None'}")
    if analysis["suggestions"]:
        out.append("  Suggestions    :")
        for s in analysis["suggestions"]:
            out.append(f"    - {s}")

    out.append(f"\n{SEP}")
    out.append("  END OF RESUME SCAN")
    out.append(SEP)
    return "\n".join(out)


# API Endpoints

@app.post("/scan")
async def scan_resume(file: UploadFile = File(...)):
    filename = file.filename or "resume"
    ext      = Path(filename).suffix.lower()

    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF or DOCX file.",
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        raw = extract_pdf(tmp_path) if ext == ".pdf" else extract_docx(tmp_path)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to read file: {str(e)}")
    finally:
        os.unlink(tmp_path)

    if not raw.strip():
        raise HTTPException(
            status_code=422,
            detail="Could not extract text from the file. It may be image-based or corrupted.",
        )

    parsed   = parse_resume(raw)
    analysis = analyze_resume(parsed)
    preview  = format_txt(parsed, analysis, filename)

    out_name = f"{Path(filename).stem}_scanned.txt"
    (OUTPUT_DIR / out_name).write_text(preview, encoding="utf-8")

    return {
        "filename": out_name,
        "parsed":   parsed,
        "analysis": analysis,
        "preview":  preview,
    }


# Ensuring only files inside outputs/ can be accessed

@app.get("/download/{filename}")
async def download(filename: str):
    # Prevent path traversal
    safe_path = (OUTPUT_DIR / Path(filename).name).resolve()
    if not str(safe_path).startswith(str(OUTPUT_DIR.resolve())):
        raise HTTPException(status_code=400, detail="Invalid filename.")
    if not safe_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")
    return FileResponse(str(safe_path), media_type="text/plain", filename=filename)


@app.get("/health")
async def health():
    return {"status": "ok"}